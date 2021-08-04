#Aplicação de automação para gerar apostilas com imagem e cabeçalho

#Imports
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import os.path

#Função que gera as paginas
def gerateExercise(teacher,year,mounth,w1_1,w1_2,w2_1,w2_2,i_exW1,i_exW2):

    #Loop para gerar as atividades
    for w in range(2):

        document = Document('doc-raiz/doc-raiz.docx')

        if w==0:
            i_ex = i_exW1
        
        if w==1:
            i_ex = i_exW2

        #Loop da primeira semana
        for i in range(i_ex):
            
            if w==0:

                if (os.path.exists(f'img/{i+1}p.jpeg')):
                    dicI = 'p'
                    dic = 'PORTUGUÊS'
                elif (os.path.exists(f'img/{i+1}m.jpeg')):
                    dicI = 'm'
                    dic = 'MATEMÁTICA'
                elif (os.path.exists(f'img/{i+1}h.jpeg')):
                    dicI = 'h'
                    dic = 'HISTÓRIA'
                elif (os.path.exists(f'img/{i+1}g.jpeg')):
                    dicI = 'g'
                    dic = 'GEOGRAFIA'
                elif (os.path.exists(f'img/{i+1}a.jpeg')):
                    dicI = 'a'
                    dic = 'ARTES'
                elif (os.path.exists(f'img/{i+1}c.jpeg')):
                    dicI = 'c'
                    dic = 'CIÊNCIAS'

                w1 = w1_1
                w2 = w1_2

                header = document.add_paragraph(f'PROFESSOR(A): {teacher}       DISCIPLINA: {dic}')
                header2 = document.add_paragraph()
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header2_text = f'ATIVIDADE PARA {year}° ANO DE ESCOLARIDADE DO MÊS DE {mounth}\nSEMANA DE {w1}  A  {w2} ATIVIDADE No:{i+1}'
                header2.add_run(header2_text).bold = True

                exercise = document.add_picture(f'img/{i+1}{dicI}.jpeg', width = Inches(7), height = Inches(8.3))

                print(w1)

            
            if w==1:

                if (os.path.exists(f'img/{i+1+i_exW1}p.jpeg')):
                    dicI = 'p'
                    dic = 'PORTUGUÊS'
                if (os.path.exists(f'img/{i+1+i_exW1}m.jpeg')):
                    dicI = 'm'
                    dic = 'MATEMÁTICA'
                if (os.path.exists(f'img/{i+1+i_exW1}h.jpeg')):
                    dicI = 'h'
                    dic = 'HISTÓRIA'
                if (os.path.exists(f'img/{i+1+i_exW1}g.jpeg')):
                    dicI = 'g'
                    dic = 'GEOGRAFIA'
                if (os.path.exists(f'img/{i+1+i_exW1}a.jpeg')):
                    dicI = 'a'
                    dic = 'ARTES'
                if (os.path.exists(f'img/{i+1+i_exW1}c.jpeg')):
                    dicI = 'c'
                    dic = 'CIÊNCIAS'

                w1 = w2_1
                w2 = w2_2

                header = document.add_paragraph(f'PROFESSOR(A): {teacher}       DISCIPLINA: {dic}')
                header2 = document.add_paragraph()
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header2_text = f'ATIVIDADE PARA {year}° ANO DE ESCOLARIDADE DO MÊS DE {mounth}\nSEMANA DE {w1}  A  {w2} ATIVIDADE No:{i+1+i_exW1}'
                header2.add_run(header2_text).bold = True
                
                exercise = document.add_picture(f'img/{i+1+i_exW1}{dicI}.jpeg', width = Inches(7), height = Inches(8.3))

                print(w2)

        document.save(f'saidas/Prof.{teacher} {mounth}{w+1}.docx')
        

#Entradas
teacher = input('Qual o nome do Professor(a)? ').upper()
year = int(input('Qual o ano de escolaridade da turma? '))
mounth = input('Qual o mês da apostila? ').upper()
i_exW1 = int(input('Quantas atividades na primeira semana? '))
i_exW2 = int(input('Quantas atividades na segunda semana? '))
w1Data1 = int(input('Que dia começa a primeira semana? '))
w1Data2 = int(input('Que dia termina a primeira semana? '))
w2Data1 = int(input('Que dia começa a segunda semana? '))
w2Data2 = int(input('Que dia termina a segunda semana? '))

#Função que trata dos dados da data das atividades

mounthNumber = 0

#Verificação do mes
if mounth == 'JANEIRO':
    mounthNumber = '01'
    
if mounth == 'FEVEREIRO':
    mounthNumber = '02'

if mounth == 'MARÇO':
    mounthNumber = '03'
    
if mounth == 'ABRIL':
    mounthNumber = '04'

if mounth == 'MAIO':
    mounthNumber = '05'
    
if mounth == 'JUNHO':
    mounthNumber = '06'

if mounth == 'JULHO':
    mounthNumber = '07'
    
if mounth == 'AGOSTO':
    mounthNumber = '08'

if mounth == 'SETEMBRO':
    mounthNumber = '09'
    
if mounth == 'OUTUBRO':
    mounthNumber = '10'

if mounth == 'NOVEMBRO':
    mounthNumber = '11'
    
if mounth == 'DEZEMBRO':
    mounthNumber = '12'

#Tratamento da data
w1_1 = f'{w1Data1}/{mounthNumber}/2021'
w2_1 = f'{w2Data1}/{mounthNumber}/2021'
w1_2 = f'{w1Data2}/{mounthNumber}/2021'
w2_2 = f'{w2Data2}/{mounthNumber}/2021'

#Onde chamamos a função que gera a apostila
gerateExercise(teacher,year,mounth,w1_1,w1_2,w2_1,w2_2,i_exW1,i_exW2)